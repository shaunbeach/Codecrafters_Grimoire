# The chancery files: a sealed writ and a bound volume.
from pypdf import PdfWriter
import docx


def reset_chancery():
    writer = PdfWriter()
    for i in range(5):
        writer.add_blank_page(width=200 + i * 10, height=280)
    with open("/workspace/writ.pdf", "wb") as handle:
        writer.write(handle)

    document = docx.Document()
    document.add_heading("Quarterly Review", level=0)
    document.add_heading("Findings", level=1)
    document.add_paragraph("Sales rose in the North.")
    document.add_paragraph("")
    document.add_heading("Recommendations", level=1)
    document.add_paragraph("Buy more rope.")
    document.add_paragraph("Ropes", style="List Bullet")
    document.save("/workspace/report.docx")


import os
os.makedirs("/workspace", exist_ok=True)
reset_chancery()
