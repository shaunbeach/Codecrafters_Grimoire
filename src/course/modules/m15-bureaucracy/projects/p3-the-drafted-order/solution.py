import docx


def draft_order(path, title, sections):
    document = docx.Document()
    document.add_heading(title, level=0)

    for name, lines in sections:
        document.add_heading(name, level=1)
        for line in lines:
            document.add_paragraph(line, style="List Bullet")

    document.save(path)
    return len(sections)
