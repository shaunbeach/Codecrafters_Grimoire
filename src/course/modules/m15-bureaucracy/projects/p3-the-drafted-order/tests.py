import docx

ORDER = "/workspace/order.docx"
SECTIONS = [
    ("Rope", ["Forty fathoms", "Delivered by Friday"]),
    ("Lanterns", ["Six, oil-fed"]),
]


def _paragraphs():
    return [(p.style.name, p.text) for p in docx.Document(ORDER).paragraphs if p.text.strip()]


def test_returns_the_section_count():
    """It reports how many sections it wrote"""
    draft_order = require("draft_order")
    assert draft_order(ORDER, "Order of Supply", SECTIONS) == 2, (
        f"Got {draft_order(ORDER, 'Order of Supply', SECTIONS)!r}"
    )


def test_the_title():
    """The title uses the Title style"""
    draft_order = require("draft_order")
    draft_order(ORDER, "Order of Supply", SECTIONS)
    assert _paragraphs()[0] == ("Title", "Order of Supply"), f"Got {_paragraphs()[0]}"


def test_sections_are_headings():
    """Section names are real Heading 1 paragraphs"""
    draft_order = require("draft_order")
    draft_order(ORDER, "Order of Supply", SECTIONS)
    headings = [text for style, text in _paragraphs() if style == "Heading 1"]
    assert headings == ["Rope", "Lanterns"], (
        f"Got {headings}. Bold text in a Normal paragraph looks like a heading and "
        "carries none of the meaning — use add_heading."
    )


def test_lines_are_bullets():
    """Each line is a List Bullet paragraph"""
    draft_order = require("draft_order")
    draft_order(ORDER, "Order of Supply", SECTIONS)
    bullets = [text for style, text in _paragraphs() if style == "List Bullet"]
    assert bullets == ["Forty fathoms", "Delivered by Friday", "Six, oil-fed"], (
        f"Got {bullets}"
    )


def test_order_is_kept():
    """Everything appears in the order it was given"""
    draft_order = require("draft_order")
    draft_order(ORDER, "Order of Supply", SECTIONS)
    texts = [text for _style, text in _paragraphs()]
    assert texts.index("Rope") < texts.index("Forty fathoms") < texts.index("Lanterns"), (
        f"Got {texts}"
    )


def test_empty_section_still_gets_a_heading():
    """A section with nothing under it is still a section"""
    draft_order = require("draft_order")
    assert draft_order(ORDER, "Short Order", [("Nothing", [])]) == 1
    headings = [text for style, text in _paragraphs() if style == "Heading 1"]
    assert headings == ["Nothing"], f"Got {headings}"
